import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_file, docx_file):
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found.")
        return

    doc = Document()
    
    # Custom styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Helvetica'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.rstrip()
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Lists (simple)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
            
        # Code blocks (very simple)
        elif line.startswith('```'):
            continue
            
        # Normal text
        elif line.strip():
            # Basic bold handling
            text = line
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph()

    doc.save(docx_file)
    print(f"Successfully saved to {docx_file}")

# Convert both documents
manifesto_md = "documentation/PLATFORM-MANIFESTO.md"
manifesto_docx = "documentation/PLATFORM-MANIFESTO.docx"
walkthrough_md = r"C:\Users\Prasad\.gemini\antigravity\brain\791c9d95-52da-44de-bff4-948cabe1c6ac\walkthrough.md"
walkthrough_docx = "documentation/PRODUCT-WALKTHROUGH.docx"

# Ensure documentation dir exists
if not os.path.exists("documentation"):
    os.makedirs("documentation")

convert_md_to_docx(manifesto_md, manifesto_docx)
convert_md_to_docx(walkthrough_md, walkthrough_docx)
